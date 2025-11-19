import streamlit as st
import re
from io import BytesIO
import fitz
from docx import Document
from docx.shared import RGBColor
import pandas as pd
from datetime import datetime
from PIL import Image, ImageDraw, ImageFont
import pytesseract
import numpy as np
import cv2

st.set_page_config(
page_title=‚ÄòAnonymiseur de Documents Medicaux‚Äô,
page_icon=‚Äòüè•‚Äô,
layout=‚Äòwide‚Äô
)

st.title(‚Äòüè• Anonymiseur de Documents Medicaux‚Äô)
st.markdown(‚Äô‚Äî‚Äô)

PATTERNS = {
‚Äòdates‚Äô: r‚Äô\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b‚Äô,
‚Äònumeros_longs‚Äô: r‚Äô\b\d{6,}\b‚Äô,
‚Äònoms_propres‚Äô: r‚Äô\b[A-Z√â√à√ä√ã√Ä√Ç√Ñ√î√ñ√õ√ú√á][a-z√©√®√™√´√†√¢√§√¥√∂√ª√º√ß]+(?:\s+[A-Z√â√à√ä√ã√Ä√Ç√Ñ√î√ñ√õ√ú√á][a-z√©√®√™√´√†√¢√§√¥√∂√ª√º√ß]+)*\b‚Äô,
‚Äòemail‚Äô: r‚Äô\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+.[A-Z|a-z]{2,}\b‚Äô,
‚Äòtelephone‚Äô: r‚Äô\b(?:+33|0)[1-9](?:[\s.-]?\d{2}){4}\b‚Äô,
‚Äònumero_secu‚Äô: r‚Äô\b[12]\s?\d{2}\s?\d{2}\s?\d{2}\s?\d{3}\s?\d{3}\s?\d{2}\b‚Äô
}

LABELS_COMMUNS = [
‚ÄòNom‚Äô, ‚ÄòPrenom‚Äô, ‚ÄòN¬∞ patient‚Äô, ‚ÄòNumero patient‚Äô, ‚ÄòPatient‚Äô,
‚ÄòAge‚Äô, ‚ÄòDate de naissance‚Äô, ‚ÄòNe(e) le‚Äô,
‚ÄòEtablissement‚Äô, ‚ÄòHopital‚Äô, ‚ÄòClinique‚Äô,
‚ÄòDate etude‚Äô, ‚ÄúDate d‚Äôetude‚Äù, ‚ÄòDate examen‚Äô,
‚ÄòEffectue par‚Äô, ‚ÄòRealise par‚Äô, ‚ÄòMedecin‚Äô, ‚ÄòDocteur‚Äô, ‚ÄòDr‚Äô,
‚ÄòAdresse‚Äô, ‚ÄòTelephone‚Äô, ‚ÄòTel‚Äô, ‚ÄòEmail‚Äô, ‚ÄòN¬∞SS‚Äô, ‚ÄòSecurite sociale‚Äô
]

def anonymize_text(text, labels_to_remove):
anonymized = text
replacements = []

```
for match in re.finditer(PATTERNS['dates'], text):
    original = match.group()
    anonymized = anonymized.replace(original, '[DATE ANONYMISEE]')
    replacements.append(('Date', original, '[DATE ANONYMISEE]'))

for match in re.finditer(PATTERNS['numeros_longs'], text):
    original = match.group()
    if not re.search(r'\d{1,2}[/-]' + re.escape(original), text):
        anonymized = anonymized.replace(original, '[NUMERO ANONYMISE]')
        replacements.append(('Numero', original, '[NUMERO ANONYMISE]'))

for match in re.finditer(PATTERNS['email'], text):
    original = match.group()
    anonymized = anonymized.replace(original, '[EMAIL ANONYMISE]')
    replacements.append(('Email', original, '[EMAIL ANONYMISE]'))

for match in re.finditer(PATTERNS['telephone'], text):
    original = match.group()
    anonymized = anonymized.replace(original, '[TEL ANONYMISE]')
    replacements.append(('Telephone', original, '[TEL ANONYMISE]'))

for match in re.finditer(PATTERNS['numero_secu'], text):
    original = match.group()
    anonymized = anonymized.replace(original, '[N¬∞SS ANONYMISE]')
    replacements.append(('N¬∞SS', original, '[N¬∞SS ANONYMISE]'))

for label in labels_to_remove:
    pattern = rf'{re.escape(label)}\s*:?\s*([^\n]+)'
    for match in re.finditer(pattern, anonymized, re.IGNORECASE):
        full_match = match.group(0)
        value = match.group(1).strip()
        if value and len(value) > 0:
            replacement = f'{label}: [ANONYMISE]'
            anonymized = anonymized.replace(full_match, replacement)
            replacements.append((label, value, '[ANONYMISE]'))

return anonymized, replacements
```

def anonymize_pdf(pdf_bytes, labels_to_remove):
doc = fitz.open(stream=pdf_bytes, filetype=‚Äòpdf‚Äô)
all_replacements = []

```
for page_num in range(len(doc)):
    page = doc[page_num]
    text = page.get_text()
    
    anonymized_text, replacements = anonymize_text(text, labels_to_remove)
    all_replacements.extend(replacements)
    
    for label in labels_to_remove:
        areas = page.search_for(label, flags=fitz.TEXT_PRESERVE_WHITESPACE)
        for area in areas:
            extended_area = fitz.Rect(area.x0, area.y0, area.x0 + 300, area.y1)
            page.add_redact_annot(extended_area, fill=(0, 0, 0))
    
    for match in re.finditer(PATTERNS['dates'], text):
        areas = page.search_for(match.group())
        for area in areas:
            page.add_redact_annot(area, fill=(0, 0, 0))
    
    for match in re.finditer(PATTERNS['numeros_longs'], text):
        areas = page.search_for(match.group())
        for area in areas:
            page.add_redact_annot(area, fill=(0, 0, 0))
    
    for match in re.finditer(PATTERNS['email'], text):
        areas = page.search_for(match.group())
        for area in areas:
            page.add_redact_annot(area, fill=(0, 0, 0))
    
    for match in re.finditer(PATTERNS['telephone'], text):
        areas = page.search_for(match.group())
        for area in areas:
            page.add_redact_annot(area, fill=(0, 0, 0))
    
    page.apply_redactions()

output_bytes = doc.write()
doc.close()

return output_bytes, all_replacements
```

def anonymize_docx(docx_bytes, labels_to_remove):
doc = Document(BytesIO(docx_bytes))
all_replacements = []

```
for para in doc.paragraphs:
    if para.text.strip():
        anonymized_text, replacements = anonymize_text(para.text, labels_to_remove)
        all_replacements.extend(replacements)
        para.text = anonymized_text

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            if cell.text.strip():
                anonymized_text, replacements = anonymize_text(cell.text, labels_to_remove)
                all_replacements.extend(replacements)
                cell.text = anonymized_text

output_buffer = BytesIO()
doc.save(output_buffer)
output_buffer.seek(0)

return output_buffer.getvalue(), all_replacements
```

def anonymize_txt(txt_bytes, labels_to_remove):
text = txt_bytes.decode(‚Äòutf-8‚Äô, errors=‚Äòignore‚Äô)
anonymized_text, replacements = anonymize_text(text, labels_to_remove)
return anonymized_text.encode(‚Äòutf-8‚Äô), replacements

def anonymize_image(image_bytes, labels_to_remove, use_ocr=True):
image = Image.open(BytesIO(image_bytes))

```
if image.mode != 'RGB':
    image = image.convert('RGB')

anonymized_image = image.copy()
draw = ImageDraw.Draw(anonymized_image)

all_replacements = []

if use_ocr:
    try:
        ocr_data = pytesseract.image_to_data(image, lang='fra+eng', output_type=pytesseract.Output.DICT)
        
        n_boxes = len(ocr_data['text'])
        for i in range(n_boxes):
            text = ocr_data['text'][i].strip()
            
            if text:
                conf = int(ocr_data['conf'][i])
                
                if conf > 30:
                    should_anonymize = False
                    replacement_type = ''
                    
                    if re.match(PATTERNS['dates'], text):
                        should_anonymize = True
                        replacement_type = 'Date'
                    
                    elif re.match(PATTERNS['numeros_longs'], text):
                        should_anonymize = True
                        replacement_type = 'Numero'
                    
                    elif re.match(PATTERNS['email'], text):
                        should_anonymize = True
                        replacement_type = 'Email'
                    
                    elif re.match(PATTERNS['telephone'], text):
                        should_anonymize = True
                        replacement_type = 'Telephone'
                    
                    else:
                        for label in labels_to_remove:
                            if label.lower() in text.lower():
                                should_anonymize = True
                                replacement_type = label
                                break
                    
                    if should_anonymize:
                        x, y, w, h = (ocr_data['left'][i], 
                                    ocr_data['top'][i], 
                                    ocr_data['width'][i], 
                                    ocr_data['height'][i])
                        
                        padding = 5
                        x -= padding
                        y -= padding
                        w += padding * 2
                        h += padding * 2
                        
                        draw.rectangle([x, y, x + w, y + h], fill='black')
                        
                        all_replacements.append((replacement_type, text, '[ANONYMISE]'))
    
    except Exception as e:
        st.warning(f'OCR non disponible ou erreur: {str(e)}. Anonymisation manuelle appliquee.')

try:
    img_array = np.array(image)
    gray = cv2.cvtColor(img_array, cv2.COLOR_RGB2GRAY)
    
    thresh = cv2.adaptiveThreshold(gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, 
                                   cv2.THRESH_BINARY, 11, 2)
    
    contours, _ = cv2.findContours(thresh, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
    
    for contour in contours:
        x, y, w, h = cv2.boundingRect(contour)
        
        if 20 < w < image.width * 0.8 and 10 < h < 100:
            roi = gray[y:y+h, x:x+w]
            
            white_pixel_ratio = np.sum(roi > 200) / (w * h)
            
            if 0.3 < white_pixel_ratio < 0.95:
                if y < image.height * 0.3:
                    draw.rectangle([x, y, x + w, y + h], fill='black')
                    all_replacements.append(('Zone detectee', f'Position ({x},{y})', '[MASQUE]'))

except Exception as e:
    st.warning(f'Detection automatique de zones limitee: {str(e)}')

output_buffer = BytesIO()
anonymized_image.save(output_buffer, format=image.format if image.format else 'PNG')
output_buffer.seek(0)

return output_buffer.getvalue(), all_replacements, image.format if image.format else 'PNG'
```

st.sidebar.header(‚ÄòConfiguration‚Äô)

st.sidebar.subheader(‚ÄòLabels a anonymiser‚Äô)
selected_labels = st.sidebar.multiselect(
‚ÄòSelectionnez les champs a anonymiser:‚Äô,
LABELS_COMMUNS,
default=[‚ÄòNom‚Äô, ‚ÄòPrenom‚Äô, ‚ÄòN¬∞ patient‚Äô, ‚ÄòAge‚Äô, ‚ÄòDate de naissance‚Äô,
‚ÄòEtablissement‚Äô, ‚ÄòDate etude‚Äô, ‚ÄòEffectue par‚Äô]
)

custom_labels = st.sidebar.text_area(
‚ÄòLabels personnalises (un par ligne):‚Äô,
help=‚ÄòAjoutez des labels supplementaires a anonymiser‚Äô
)

if custom_labels:
custom_labels_list = [label.strip() for label in custom_labels.split(‚Äô\n‚Äô) if label.strip()]
selected_labels.extend(custom_labels_list)

st.sidebar.subheader(‚ÄòOptions pour les images‚Äô)
use_ocr = st.sidebar.checkbox(
‚ÄúUtiliser l‚ÄôOCR (reconnaissance de texte)‚Äù,
value=True,
help=‚ÄòActive la detection automatique de texte dans les images‚Äô
)

st.sidebar.markdown(‚Äô‚Äî‚Äô)
st.sidebar.info(
‚ÄòInformation\n\n‚Äô
‚ÄòCette application anonymise automatiquement:\n‚Äô
‚Äò- Les dates (JJ/MM/AAAA)\n‚Äô
‚Äò- Les numeros longs (6+ chiffres)\n‚Äô
‚Äò- Les emails\n‚Äô
‚Äò- Les numeros de telephone\n‚Äô
‚Äò- Les numeros de securite sociale\n‚Äô
‚Äò- Les champs selectionnes\n‚Äô
‚Äú- Le texte dans les images (OCR)‚Äù
)

st.subheader(‚ÄòCharger le document medical‚Äô)
uploaded_file = st.file_uploader(
‚ÄòChoisissez un fichier (PDF, Word, TXT ou Image)‚Äô,
type=[‚Äòpdf‚Äô, ‚Äòdocx‚Äô, ‚Äòdoc‚Äô, ‚Äòtxt‚Äô, ‚Äòpng‚Äô, ‚Äòjpg‚Äô, ‚Äòjpeg‚Äô, ‚Äògif‚Äô, ‚Äòbmp‚Äô, ‚Äòtiff‚Äô],
help=‚ÄòFormats acceptes: PDF, DOCX, TXT, PNG, JPG, JPEG, GIF, BMP, TIFF‚Äô
)

if uploaded_file is not None:
st.success(f‚ÄôFichier charge: {uploaded_file.name}‚Äô)

```
file_extension = uploaded_file.name.split('.')[-1].lower()
if file_extension in ['png', 'jpg', 'jpeg', 'gif', 'bmp', 'tiff']:
    col1, col2 = st.columns(2)
    with col1:
        st.subheader('Image originale')
        st.image(uploaded_file, use_container_width=True)

if st.button('Anonymiser le document', type='primary'):
    with st.spinner('Anonymisation en cours...'):
        try:
            file_bytes = uploaded_file.read()
            file_extension = uploaded_file.name.split('.')[-1].lower()
            
            if file_extension == 'pdf':
                anonymized_bytes, replacements = anonymize_pdf(file_bytes, selected_labels)
                mime_type = 'application/pdf'
                output_extension = 'pdf'
                
            elif file_extension in ['docx', 'doc']:
                anonymized_bytes, replacements = anonymize_docx(file_bytes, selected_labels)
                mime_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                output_extension = 'docx'
                
            elif file_extension == 'txt':
                anonymized_bytes, replacements = anonymize_txt(file_bytes, selected_labels)
                mime_type = 'text/plain'
                output_extension = 'txt'
            
            elif file_extension in ['png', 'jpg', 'jpeg', 'gif', 'bmp', 'tiff']:
                anonymized_bytes, replacements, img_format = anonymize_image(
                    file_bytes, selected_labels, use_ocr
                )
                mime_type = f'image/{img_format.lower()}'
                output_extension = img_format.lower()
            
            st.success('Anonymisation terminee!')
            
            if file_extension in ['png', 'jpg', 'jpeg', 'gif', 'bmp', 'tiff']:
                with col2:
                    st.subheader('Image anonymisee')
                    st.image(anonymized_bytes, use_container_width=True)
            
            col_stat1, col_stat2 = st.columns(2)
            with col_stat1:
                st.metric('Elements anonymises', len(replacements))
            with col_stat2:
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            
            if replacements:
                st.subheader('Details des anonymisations')
                df_replacements = pd.DataFrame(
                    replacements,
                    columns=['Type', 'Valeur originale', 'Remplacement']
                )
                st.dataframe(df_replacements, use_container_width=True)
            else:
                st.info('Aucune donnee sensible detectee automatiquement.')
            
            st.subheader('Telecharger le document anonymise')
            original_name = uploaded_file.name.rsplit('.', 1)[0]
            output_filename = f'{original_name}_anonymise_{timestamp}.{output_extension}'
            
            st.download_button(
                label=f'Telecharger {output_filename}',
                data=anonymized_bytes,
                file_name=output_filename,
                mime=mime_type,
                type='primary'
            )
            
            st.warning(
                'Attention: Verifiez toujours manuellement le document anonymise '
                'avant de le partager pour vous assurer que toutes les donnees sensibles '
                'ont ete correctement supprimees.'
            )
            
        except Exception as e:
            st.error(f'Erreur lors de l\'anonymisation: {str(e)}')
            st.exception(e)
```

else:
st.info(
‚ÄòPour commencer:\n\n‚Äô
‚Äò1. Selectionnez les champs a anonymiser dans la barre laterale\n‚Äô
‚Äò2. Telechargez votre document medical (PDF, Word, TXT ou Image)\n‚Äô
‚Äú3. Cliquez sur ‚ÄòAnonymiser le document‚Äô\n‚Äù
‚Äò4. Telechargez le document anonymise‚Äô
)

```
with st.expander('Types de fichiers supportes'):
    st.markdown('''
    **Documents texte:**
    - PDF (avec masquage visuel des donnees)
    - Word (.docx)
    - Fichiers texte (.txt)
    
    **Images medicales:**
    - PNG
    - JPG / JPEG
    - GIF
    - BMP
    - TIFF
    
    Pour les images, l'OCR detecte automatiquement le texte et masque:
    - Les informations d'en-tete (nom, date, numero)
    - Les dates et numeros dans l'image
    - Les zones de texte personnalisees
    ''')
```

st.markdown(‚Äô‚Äî‚Äô)
st.markdown(
‚Äú<div style='text-align: center; color: gray;'>‚Äù
‚ÄúApplication d‚Äôanonymisation de documents medicaux | ‚Äú
‚ÄôDeveloppe pour la protection des donnees patients | ‚Äô
‚ÄòSupport: PDF, Word, TXT, Images‚Äô
‚Äò</div>‚Äô,
unsafe_allow_html=True
)
